import Tkinter
from tkFileDialog import askopenfilename, askdirectory      
import re
from docx import Document

class vocass(Tkinter.Tk):
    def __init__(self,parent):
        Tkinter.Tk.__init__(self,parent)
        self.parent = parent
        self.initialize()

    def initialize(self):
        self.grid()

        frame_wl = Tkinter.Frame(self,bd=2,width=600,height=40,relief="sunken")
        frame_wl.pack(fill="both",expand=1)
        frame_wl.pack_propagate(0)
        frame_doc = Tkinter.Frame(self,bd=2,width=600,height=40,relief="sunken")
        frame_doc.pack(fill="both",expand=1)
        frame_doc.pack_propagate(0)
        frame_output = Tkinter.Frame(self,bd=2,width=600,height=40,relief="sunken")
        frame_output.pack(fill="both",expand=1)
        frame_output.pack_propagate(0)
        frame_run = Tkinter.Frame(self,bd=2,width=600,height=40,relief="sunken")
        frame_run.pack(fill="both",expand=1)
        frame_run.pack_propagate(0)

        button_wl = Tkinter.Button(frame_wl,width=20,text=u"Select A Word List",command=self.selectWordList)
        button_wl.pack(side="left")

        self.wordlistselected = Tkinter.StringVar()
        label_wl = Tkinter.Label(frame_wl,textvariable=self.wordlistselected,anchor="w",fg="red")
        self.wordlistselected.set("No File Selected")
        label_wl.pack(side="left")

        button_doc = Tkinter.Button(frame_doc,width=20,text=u"Select A Document",command=self.selectDocument)
        button_doc.pack(side="left")

        self.documentselected = Tkinter.StringVar()
        label_doc = Tkinter.Label(frame_doc,textvariable=self.documentselected,anchor="w",fg="red")
        self.documentselected.set("No File Selected")
        label_doc.pack(side="left")

        button_output = Tkinter.Button(frame_output,width=20,text=u"Select Output Directory",command=self.selectDirectory)
        button_output.pack(side="left")

        self.directoryselected = Tkinter.StringVar()
        label_output = Tkinter.Label(frame_output,textvariable=self.directoryselected,anchor="w",fg="red")
        self.directoryselected.set("No Directory Selected")
        label_output.pack(side="left")

        button_run = Tkinter.Button(frame_run,width=20,text=u"Run!",command=self.run)
        button_run.pack(side="left")

        self.message = Tkinter.StringVar()
        label_run = Tkinter.Label(frame_run,textvariable=self.message,anchor="w",fg="red")
        self.message.set("Nothing Generated")
        label_run.pack(side="left")       

    def selectWordList(self):
        filename = askopenfilename(filetypes=[("Text Files", "*.txt")])
        if (len(filename) >= 1):
        	self.wordlistselected.set(filename)

    def selectDocument(self):
        filename = askopenfilename(filetypes=[("Text Files", "*.txt")])
        if (len(filename) >= 1):
        	self.documentselected.set(filename)

    def selectDirectory(self):
    	directoryname = askdirectory()
    	if (len(directoryname) >= 1):
        	self.directoryselected.set(directoryname)

    def run(self):
    	f = open(self.wordlistselected.get())
    	voc = []
    	for line in f:
    		for word in line.split():
    			voc.append(word)
        f.close()

    	f = open(self.documentselected.get())
        dic = {}
        text = f.read().lower()
        # content = f.read()
        # text = re.sub('[,.!?*\(\)\"\'|#&$\[\];\{\}_:/]', " ", text)
        text = re.sub('\. ', " ", text)
        text = re.sub('\.\)', " ", text)
        text = re.sub('\.\"', " ", text)
        text = re.sub('\"\'', " ", text)
        text = re.sub('\.\'', " ", text)
        text = re.sub(' \'', " ", text)
        text = re.sub('\' ', " ", text)
        text = re.sub('\.\\n', " ", text)
        text = re.sub('\.\\r\\n', " ", text)
        text = re.sub('--', " ", text)
        text = re.sub('[^a-z\'\-\.]+', " ", text)
        # text = text.translate(None, string.punctuation)
        for word in text.split():
            if ((word in voc) == False):
                i = dic.get(word)
                if (i == None):
                    dic[word] = 1
                else:
                    dic[word] = dic[word] + 1
        f.close()

    	document = Document()
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Word'
        hdr_cells[1].text = '# of times'
        hdr_cells[2].text = 'Difficulty level'

    	for word in sorted(dic):
            row_cells = table.add_row().cells
            row_cells[0].text = str(word)
            row_cells[1].text = str(dic[word])

        document.save(self.directoryselected.get() + '/output.docx')

    	self.message.set("output.docx Successfully Generated!")

if __name__ == "__main__":
    application = vocass(None)
    application.title('Vocabulary Assessor')
    application.mainloop()